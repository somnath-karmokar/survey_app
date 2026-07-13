#!/usr/bin/env python
"""
Standalone script to import Journal posts from .docx files.

Reads every .docx file in the `Journals/` directory and creates a
JournalPost for each one:
  - the first "Heading 2" paragraph becomes the title
  - subsequent "Heading 2" paragraphs become <h2> section headings
  - normal paragraphs become <p> (bold/italic runs preserved)
  - tables become HTML <table>s
  - the excerpt is auto-derived from the first body paragraph

Usage (run from the project root, with the venv active):
    python upload_journals.py                 # import new posts
    python upload_journals.py --dry-run        # preview only, no DB writes
    python upload_journals.py --update         # overwrite posts that already exist (matched by title)
    python upload_journals.py --dir "Journals" # custom source directory
    python upload_journals.py --unpublished    # create posts with is_published=False
"""
import argparse
import html
import os
import sys

import django

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'survey_app.settings')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
django.setup()

from docx import Document  # noqa: E402  (must come after django.setup())

from surveys.models import JournalPost  # noqa: E402

EXCERPT_MAX_LEN = 300
HEADING_MAX_WORDS = 12


def run_to_html(run):
    text = html.escape(run.text)
    if not text:
        return ''
    if run.bold:
        text = f'<strong>{text}</strong>'
    if run.italic:
        text = f'<em>{text}</em>'
    if run.underline:
        text = f'<u>{text}</u>'
    return text


def paragraph_to_html(paragraph, tag='p'):
    inner = ''.join(run_to_html(r) for r in paragraph.runs).strip()
    if not inner:
        return ''
    return f'<{tag}>{inner}</{tag}>'


def is_heading_paragraph(paragraph):
    """
    The source .docx files mostly use no real Word heading style at all --
    every section title is just a short plain paragraph. Detect them by shape
    instead: a single run, no trailing sentence punctuation, not a bullet
    ("**Term**: description"), not a decorative/ASCII-art line, and short.
    """
    text = paragraph.text.strip()
    if not text or not any(c.isalpha() for c in text):
        return False

    if paragraph.style.name.startswith('Heading'):
        return True

    runs = paragraph.runs
    is_bullet = len(runs) >= 2 and bool(runs[0].bold)
    has_arrow = '--->' in text
    ends_with_sentence_punct = text.endswith(('.', '!', '?'))
    word_count = len(text.split())

    return (
        not is_bullet
        and not has_arrow
        and not ends_with_sentence_punct
        and word_count <= HEADING_MAX_WORDS
    )


def table_to_html(table):
    rows_html = []
    for i, row in enumerate(table.rows):
        cell_tag = 'th' if i == 0 else 'td'
        cells = ''.join(
            f'<{cell_tag}>{html.escape(cell.text.strip())}</{cell_tag}>'
            for cell in row.cells
        )
        rows_html.append(f'<tr>{cells}</tr>')
    return '<table border="1" cellpadding="6" cellspacing="0">' + ''.join(rows_html) + '</table>'


def iter_body_blocks(document):
    """Yield ('paragraph', paragraph) / ('table', table) in document order."""
    body = document.element.body
    paragraphs = {p._p: p for p in document.paragraphs}
    tables = {t._tbl: t for t in document.tables}
    for child in body:
        if child.tag.endswith('}p') and child in paragraphs:
            yield 'paragraph', paragraphs[child]
        elif child.tag.endswith('}tbl') and child in tables:
            yield 'table', tables[child]


def parse_docx(path):
    """Return (title, excerpt, content_html) or None if the file isn't a well-formed journal post.

    Every genuine journal post in this batch ends with a summary table; files
    without one (e.g. leftover scratch/notes docs) are treated as not importable.
    """
    document = Document(path)

    if not document.tables:
        return None

    title = None
    excerpt = ''
    content_parts = []

    for kind, block in iter_body_blocks(document):
        if kind == 'table':
            content_parts.append(table_to_html(block))
            continue

        paragraph = block
        text = paragraph.text.strip()
        if not text:
            continue

        if title is None:
            title = text
            continue

        if is_heading_paragraph(paragraph):
            content_parts.append(f'<h2>{html.escape(text)}</h2>')
            continue

        piece = paragraph_to_html(paragraph, tag='p')
        if not piece:
            continue
        content_parts.append(piece)
        if not excerpt:
            excerpt = text[:EXCERPT_MAX_LEN]

    if title is None:
        return None

    return title, excerpt, '\n'.join(content_parts)


def import_journals(source_dir, dry_run=False, update=False, publish=True):
    created, updated, skipped = 0, 0, 0

    filenames = sorted(f for f in os.listdir(source_dir) if f.lower().endswith('.docx'))
    if not filenames:
        print(f'No .docx files found in "{source_dir}".')
        return

    for filename in filenames:
        path = os.path.join(source_dir, filename)
        parsed = parse_docx(path)
        if parsed is None:
            print(f'[skip]    {filename} -- no summary table found (not a well-formed journal post)')
            skipped += 1
            continue

        title, excerpt, content = parsed
        existing = JournalPost.objects.filter(title=title).first()

        if existing and not update:
            print(f'[exists]  {filename} -> "{title}" (use --update to overwrite)')
            skipped += 1
            continue

        if dry_run:
            action = 'update' if existing else 'create'
            print(f'[dry-run] would {action}: "{title}" ({len(content)} chars content)')
            continue

        if existing:
            existing.excerpt = excerpt
            existing.content = content
            existing.is_published = publish
            existing.save()
            print(f'[updated] {filename} -> "{title}"')
            updated += 1
        else:
            JournalPost.objects.create(
                title=title,
                excerpt=excerpt,
                content=content,
                is_published=publish,
            )
            print(f'[created] {filename} -> "{title}"')
            created += 1

    print(f'\nDone. created={created} updated={updated} skipped={skipped}')


def main():
    parser = argparse.ArgumentParser(description='Import Journal posts from .docx files.')
    parser.add_argument(
        '--dir',
        default=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Journals'),
        help='Directory containing the .docx journal files (default: ./Journals)',
    )
    parser.add_argument('--dry-run', action='store_true', help='Preview without writing to the database.')
    parser.add_argument('--update', action='store_true', help='Overwrite posts that already exist (matched by title).')
    parser.add_argument('--unpublished', action='store_true', help='Create posts with is_published=False.')
    args = parser.parse_args()

    if not os.path.isdir(args.dir):
        print(f'Directory not found: {args.dir}')
        sys.exit(1)

    import_journals(
        args.dir,
        dry_run=args.dry_run,
        update=args.update,
        publish=not args.unpublished,
    )


if __name__ == '__main__':
    main()
